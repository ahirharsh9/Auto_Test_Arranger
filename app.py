# ★ Murlidhar Academy MCQ Generator (Final Stable Version) ★

import streamlit as st
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import tempfile
import requests
import os

# -------------------------------------------------
# CONFIG
# -------------------------------------------------

st.set_page_config(page_title="Murlidhar MCQ Generator", layout="wide")

st.title("📄 Murlidhar Academy MCQ Paper Generator (Final Stable Math)")
st.markdown("Default Template: Google Drive | Optional: Upload New Template")

DEFAULT_TEMPLATE_URL = "https://docs.google.com/document/d/1JMow6oJ2ASJah5vM4OK1Q-uYPefiMnEg/export?format=docx"

# -------------------------------------------------
# FINAL SAFE MATH FORMATTER
# -------------------------------------------------

def fix_math_formatting(text):

    if not text:
        return ""

    text = text.replace("$", "")

    # Fractions
    text = re.sub(r'\\frac\{([^{}]+)\}\{([^{}]+)\}', r'\1/\2', text)

    # Roots
    text = re.sub(r'\\sqrt\[([^\]]+)\]\{([^{}]+)\}', r'\1√(\2)', text)
    text = re.sub(r'\\sqrt\{([^{}]+)\}', r'√(\1)', text)

    # Symbols
    text = text.replace("\\times", "×")
    text = text.replace("\\div", "÷")
    text = text.replace("\\infty", "∞")

    # Superscript map
    superscripts = {
        "0":"⁰","1":"¹","2":"²","3":"³","4":"⁴",
        "5":"⁵","6":"⁶","7":"⁷","8":"⁸","9":"⁹",
        "-":"⁻","+":"⁺",".":"·",
        "a":"ᵃ","b":"ᵇ","c":"ᶜ",
        "x":"ˣ","y":"ʸ","z":"ᶻ","n":"ⁿ"
    }

    def convert_to_superscript(content):
        return "".join(superscripts.get(ch, ch) for ch in content)

    # Handle ^{...}
    text = re.sub(
        r'\^\{([^{}]+)\}',
        lambda m: convert_to_superscript(m.group(1)),
        text
    )

    # Handle ^n , ^2 , ^n-1 , ^n+1
    text = re.sub(
        r'\^([a-zA-Z0-9\+\-\.]+)',
        lambda m: convert_to_superscript(m.group(1)),
        text
    )

    text = text.replace("\\", "")

    return text


# -------------------------------------------------
# TEXT CLEANER
# -------------------------------------------------

def clean_garbage_text(text, keep_pipe=False):

    if not text:
        return ""

    text = re.sub(r'\[cite.*?\]', '', text)
    text = re.sub(r'\[source.*?\]', '', text)
    text = re.sub(r'\[cite_start\]', '', text)

    if not keep_pipe:
        text = text.replace('|', '')

    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n+', '\n', text)

    text = fix_math_formatting(text)

    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        clean_line = re.sub(r'\s+', ' ', line).strip()
        if clean_line:
            cleaned_lines.append(clean_line)

    return "\n".join(cleaned_lines)


# -------------------------------------------------
# PARSE MCQ
# -------------------------------------------------

def parse_mcq_text(raw_text):

    question_pattern = re.compile(r'(?=\(\d{2}\))')
    blocks = question_pattern.split(raw_text)

    parsed_questions = []

    for block in blocks:

        if not block.strip():
            continue

        q_num_match = re.search(r'\((\d{2})\)', block)
        if not q_num_match:
            continue

        parts = re.split(r'\(A\)', block, maxsplit=1)

        raw_q_text = parts[0].replace(q_num_match.group(0), '')
        q_text_clean = clean_garbage_text(raw_q_text, keep_pipe=True)

        options_part = ""
        if len(parts) > 1:
            options_part = "(A) " + parts[1]

        def extract_option(label1, label2=None):
            if label2:
                pattern = rf'\({label1}\)(.*?)\({label2}\)'
            else:
                pattern = rf'\({label1}\)(.*)'
            match = re.search(pattern, options_part, re.DOTALL)
            return clean_garbage_text(match.group(1), False) if match else ""

        parsed_questions.append({
            "q_num": q_num_match.group(1),
            "question": q_text_clean,
            "A": extract_option("A", "B"),
            "B": extract_option("B", "C"),
            "C": extract_option("C", "D"),
            "D": extract_option("D", None)
        })

    return parsed_questions


# -------------------------------------------------
# CREATE DOC
# -------------------------------------------------

def create_doc(template_path, questions_data):

    try:
        doc = Document(template_path)
    except:
        doc = Document()

    def set_font(run):
        font_name = 'HindVadodara'
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(11)
        run.font.bold = False

    def add_text(paragraph, text):
        parts = text.split('\n')
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            set_font(run)
            if i < len(parts) - 1:
                paragraph.add_run().add_break()

    for q in questions_data:

        p = doc.add_paragraph()

        run_num = p.add_run(f"({q['q_num']}) ")
        set_font(run_num)
        add_text(p, q['question'])
        p.add_run().add_break()

        for label in ["A", "B"]:
            p.add_run(f"({label}) ")
            set_font(p.runs[-1])
            add_text(p, q[label])
            p.add_run("\t")
            set_font(p.runs[-1])

        p.add_run().add_break()

        for label in ["C", "D"]:
            p.add_run(f"({label}) ")
            set_font(p.runs[-1])
            add_text(p, q[label])
            if label == "C":
                p.add_run("\t")
                set_font(p.runs[-1])

    output_filename = "Murlidhar_Final_Pro.docx"
    doc.save(output_filename)
    return output_filename


# -------------------------------------------------
# UI SECTION
# -------------------------------------------------

st.subheader("📂 Optional: Upload New Word Template")
uploaded_template = st.file_uploader("Upload .docx file (optional)", type=["docx"])

mcq_text = st.text_area("✍️ Paste Raw MCQs Text", height=300)

if st.button("🚀 Generate Paper"):

    if not mcq_text.strip():
        st.error("❌ Please paste MCQs.")
        st.stop()

    try:

        if uploaded_template is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(uploaded_template.read())
                template_path = tmp.name
        else:
            response = requests.get(DEFAULT_TEMPLATE_URL)
            if response.status_code != 200:
                st.error("❌ Failed to download default template.")
                st.stop()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(response.content)
                template_path = tmp.name

        q_data = parse_mcq_text(mcq_text)
        final_file = create_doc(template_path, q_data)

        st.success(f"✅ {len(q_data)} Questions Processed Successfully!")

        with open(final_file, "rb") as f:
            st.download_button(
                "📥 Download Final Paper",
                f,
                file_name="Murlidhar_Final_Pro.docx"
            )

        os.remove(template_path)

    except Exception as e:
        st.error(f"❌ Error Occurred: {e}")
